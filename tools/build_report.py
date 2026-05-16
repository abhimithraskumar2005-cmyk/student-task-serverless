from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_TITLE = "Cloud-Based Serverless Student Task Management System"
AUTHOR = "Abhimithra S Kumar"
REPORT_PATH = "Cloud_Based_Serverless_Student_Task_Management_System_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            if row_index == 0:
                set_cell_shading(cell, "E8F3F2")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(18, 52, 59)
    return heading


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Insert Screenshot: {caption}]")
    run.italic = True
    run.font.color.rgb = RGBColor(90, 90, 90)
    run.font.size = Pt(10)


def add_metadata_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.autofit = True
    rows = [
        ("Project Title", PROJECT_TITLE),
        ("Student Name", AUTHOR),
        ("Project Type", "Cloud Computing / Serverless Web Application"),
        ("Cloud Provider", "Amazon Web Services (AWS)"),
        ("Submission Date", "May 31, 2026"),
    ]
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.cell(index, 0), label, bold=True)
        set_cell_text(table.cell(index, 1), value)
    style_table(table)


def add_service_table(doc):
    rows = [
        ("Amazon S3", "Hosts the static frontend files: HTML, CSS, and JavaScript."),
        ("Amazon CloudFront", "Delivers the S3-hosted frontend through a CDN with HTTPS."),
        ("Amazon API Gateway", "Provides HTTP API routes for frontend-backend communication."),
        ("AWS Lambda", "Runs backend logic for create, read, update, and delete operations."),
        ("Amazon DynamoDB", "Stores task data using userId and taskId as the primary key."),
        ("AWS IAM", "Provides least privilege access for Lambda to use DynamoDB."),
        ("Amazon CloudWatch", "Stores logs and metrics for monitoring and debugging."),
        ("GitHub", "Stores project source code and documentation."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.cell(0, 0), "Service / Tool", bold=True)
    set_cell_text(table.cell(0, 1), "Purpose")
    for service, purpose in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], service, bold=True)
        set_cell_text(cells[1], purpose)
    style_table(table)


def add_api_table(doc):
    rows = [
        ("POST", "/tasks", "Creates a new task item in DynamoDB."),
        ("GET", "/tasks?userId=demo-student", "Fetches all tasks for a user."),
        ("PUT", "/tasks/{taskId}", "Updates an existing task."),
        ("DELETE", "/tasks/{taskId}", "Deletes a task from DynamoDB."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.cell(0, 0), "Method", bold=True)
    set_cell_text(table.cell(0, 1), "Endpoint", bold=True)
    set_cell_text(table.cell(0, 2), "Description", bold=True)
    for method, endpoint, description in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], method, bold=True)
        set_cell_text(cells[1], endpoint)
        set_cell_text(cells[2], description)
    style_table(table)


def add_testing_table(doc):
    rows = [
        ("DynamoDB table creation", "StudentTasks table becomes Active", "Pass"),
        ("Lambda POST test", "Task is created with statusCode 201", "Pass"),
        ("Lambda GET test", "Tasks returned with statusCode 200", "Pass"),
        ("API Gateway GET route", "Tasks returned through public API URL", "Pass"),
        ("Frontend local test", "Tasks load in AWS API Mode", "Pass"),
        ("S3 hosted frontend", "Frontend opens through S3 website endpoint", "Pass"),
        ("CloudFront hosted frontend", "Frontend opens through CloudFront URL", "Pass"),
        ("Monitoring", "CloudWatch logs and metrics are visible", "Pass"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.cell(0, 0), "Test Case", bold=True)
    set_cell_text(table.cell(0, 1), "Expected Result", bold=True)
    set_cell_text(table.cell(0, 2), "Status", bold=True)
    for case, expected, status in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], case)
        set_cell_text(cells[1], expected)
        set_cell_text(cells[2], status, bold=True)
    style_table(table)


def add_screenshot_table(doc):
    screenshots = [
        ("01", "DynamoDB table list active"),
        ("05", "Lambda function overview"),
        ("06", "Lambda code deployed"),
        ("07", "Lambda environment variable"),
        ("09", "IAM DynamoDB policy"),
        ("10", "Lambda create task success"),
        ("11", "Lambda get tasks success"),
        ("12", "DynamoDB task item created"),
        ("17", "API Gateway CORS settings"),
        ("18", "API Gateway stage invoke URL"),
        ("19", "API Gateway GET test success"),
        ("21", "Frontend AWS API mode"),
        ("24", "DynamoDB frontend task created"),
        ("28", "S3 static website hosting enabled"),
        ("29", "S3 bucket policy public read"),
        ("31", "S3 hosted frontend working"),
        ("40", "CloudFront hosted frontend working"),
        ("45", "CloudWatch Lambda log events"),
        ("46", "Lambda monitor metrics"),
        ("48", "CloudFront monitoring metrics"),
        ("49", "DynamoDB monitoring metrics"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.cell(0, 0), "Screenshot No.", bold=True)
    set_cell_text(table.cell(0, 1), "Evidence")
    for number, evidence in screenshots:
        cells = table.add_row().cells
        set_cell_text(cells[0], number, bold=True)
        set_cell_text(cells[1], evidence)
    style_table(table)


def add_architecture_diagram(doc):
    diagram_lines = [
        "User Browser",
        "  -> CloudFront",
        "  -> S3 Static Website Frontend",
        "  -> API Gateway",
        "  -> AWS Lambda",
        "  -> DynamoDB",
        "",
        "IAM secures Lambda permissions.",
        "CloudWatch monitors logs and metrics.",
    ]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("\n".join(diagram_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(18, 52, 59)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(PROJECT_TITLE)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(18, 52, 59)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Final Capstone Project Report")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(53, 109, 111)

    add_metadata_table(doc)
    doc.add_paragraph()

    add_heading(doc, "Abstract", 1)
    doc.add_paragraph(
        "This project presents a cloud-based serverless student task management system. "
        "The application allows students to create, view, update, and delete project tasks. "
        "The frontend is hosted using Amazon S3 and distributed through Amazon CloudFront. "
        "API Gateway exposes HTTP endpoints that invoke AWS Lambda functions, and DynamoDB stores task data. "
        "IAM is used for secure permissions, while CloudWatch provides logs and monitoring."
    )

    add_heading(doc, "Problem Statement", 1)
    doc.add_paragraph(
        "Students need a simple and reliable system to track internship and capstone project tasks. "
        "Traditional server-based systems require infrastructure management, server maintenance, and scaling effort. "
        "This project solves the problem by using a serverless AWS architecture that is scalable, cost-efficient, and easy to maintain."
    )

    add_heading(doc, "Objectives", 1)
    add_bullets(
        doc,
        [
            "Build a cloud-based task management web application.",
            "Use serverless AWS services for backend processing.",
            "Store task data in a managed NoSQL database.",
            "Host the frontend using S3 and CloudFront.",
            "Expose backend functionality using API Gateway and Lambda.",
            "Apply IAM permissions using the least privilege principle.",
            "Use CloudWatch for logs, metrics, and monitoring evidence.",
        ],
    )

    add_heading(doc, "Tools And Technologies", 1)
    add_service_table(doc)

    add_heading(doc, "System Architecture", 1)
    doc.add_paragraph(
        "The project follows a serverless architecture where the frontend, API layer, backend logic, database, and monitoring are separated into managed AWS services."
    )
    add_architecture_diagram(doc)
    add_placeholder(doc, "Architecture diagram / AWS service flow")

    add_heading(doc, "Features", 1)
    add_bullets(
        doc,
        [
            "Add new student project tasks.",
            "View task list from DynamoDB.",
            "Update task status and task details.",
            "Delete unwanted tasks.",
            "Dashboard counters for total, pending, in-progress, and completed tasks.",
            "Frontend deployed through S3 and CloudFront.",
            "Backend deployed using API Gateway and Lambda.",
        ],
    )

    add_heading(doc, "Implementation", 1)
    add_heading(doc, "Frontend", 2)
    doc.add_paragraph(
        "The frontend is built using HTML, CSS, and JavaScript. It displays a dashboard, task form, and task list. "
        "The JavaScript file is configured with the API Gateway base URL so that all task operations are sent to the AWS backend."
    )
    doc.add_paragraph("Frontend API Gateway URL: https://y93vfcqzw1.execute-api.ap-south-1.amazonaws.com")
    doc.add_paragraph("CloudFront URL: https://d14hq1kn5kv9jt.cloudfront.net")
    add_placeholder(doc, "Frontend in AWS API Mode")

    add_heading(doc, "Database", 2)
    doc.add_paragraph(
        "Amazon DynamoDB is used as the NoSQL database. The table StudentTasks uses a composite primary key with userId as the partition key and taskId as the sort key. "
        "This design allows tasks for a specific user to be queried efficiently."
    )
    add_bullets(doc, ["Table name: StudentTasks", "Partition key: userId", "Sort key: taskId", "Billing mode: On-demand"])
    add_placeholder(doc, "DynamoDB table and created task items")

    add_heading(doc, "Backend", 2)
    doc.add_paragraph(
        "AWS Lambda is used for backend logic. The Lambda function student-task-api receives requests, validates input, and performs CRUD operations on DynamoDB. "
        "The TABLE_NAME environment variable is used to connect the function to the StudentTasks table."
    )
    add_placeholder(doc, "Lambda function overview, code, and tests")

    add_heading(doc, "API Design", 2)
    doc.add_paragraph("Amazon API Gateway exposes HTTP routes that invoke the Lambda backend.")
    add_api_table(doc)

    add_heading(doc, "Deployment", 1)
    add_numbered(
        doc,
        [
            "Created DynamoDB table StudentTasks.",
            "Created Lambda function student-task-api.",
            "Added IAM policy for DynamoDB access.",
            "Created API Gateway HTTP API and configured routes.",
            "Connected frontend to API Gateway.",
            "Uploaded frontend files to S3.",
            "Enabled S3 static website hosting.",
            "Created CloudFront distribution for CDN and HTTPS access.",
        ],
    )
    add_placeholder(doc, "S3 website hosting and CloudFront working frontend")

    add_heading(doc, "Security", 1)
    doc.add_paragraph(
        "The project uses IAM roles to control access between services. The Lambda execution role is granted permission only to access the StudentTasks DynamoDB table. "
        "CloudFront provides HTTPS access for the frontend. API Gateway CORS settings are configured to allow browser-based requests from the frontend."
    )

    add_heading(doc, "Monitoring", 1)
    doc.add_paragraph(
        "CloudWatch is used to monitor the application. Lambda logs are available in the /aws/lambda/student-task-api log group. "
        "Lambda, CloudFront, and DynamoDB metrics were reviewed to confirm invocations, requests, duration, and read usage."
    )
    add_placeholder(doc, "CloudWatch Lambda logs and metrics")

    add_heading(doc, "Testing", 1)
    add_testing_table(doc)

    add_heading(doc, "Screenshot Evidence Checklist", 1)
    add_screenshot_table(doc)

    add_heading(doc, "Challenges Faced", 1)
    add_bullets(
        doc,
        [
            "Configured Lambda handler correctly after an initial handler mismatch.",
            "Added correct IAM permissions so Lambda could access DynamoDB.",
            "Configured S3 bucket policy and Block Public Access settings for static website hosting.",
            "Used the S3 website endpoint as the CloudFront origin.",
        ],
    )

    add_heading(doc, "Future Enhancements", 1)
    add_bullets(
        doc,
        [
            "Add Amazon Cognito for user signup and login.",
            "Add custom domain using Route 53 and ACM certificate.",
            "Enable API Gateway access logging.",
            "Add automated deployment using GitHub Actions.",
            "Improve UI with search, filters, and due-date notifications.",
        ],
    )

    add_heading(doc, "Conclusion", 1)
    doc.add_paragraph(
        "The project successfully demonstrates a complete AWS serverless web application. "
        "The frontend is hosted on S3 and delivered through CloudFront, while API Gateway, Lambda, and DynamoDB provide backend functionality. "
        "IAM and CloudWatch strengthen the project by adding security and monitoring. "
        "This architecture is scalable, cost-effective, and suitable for modern cloud-based applications."
    )

    add_heading(doc, "References", 1)
    add_bullets(
        doc,
        [
            "Amazon S3 Static Website Hosting Documentation",
            "Amazon CloudFront Developer Guide",
            "AWS Lambda Developer Guide",
            "Amazon API Gateway HTTP API Documentation",
            "Amazon DynamoDB Developer Guide",
            "Amazon CloudWatch Documentation",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("Cloud-Based Serverless Student Task Management System")
        footer_run.font.name = "Arial"
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(90, 90, 90)

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build()
